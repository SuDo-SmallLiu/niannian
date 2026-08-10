#!/usr/bin/env python3
"""生成念念年年完整开发文档 Word 版"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "念念年年-开发文档全集.docx"

DOCS = [
    ROOT / "docs/README.md",
    ROOT / "docs/plans/2026-08-06-current-roadmap.md",
    ROOT / "docs/DEVELOPMENT-HISTORY.md",
    ROOT / "docs/ARCHITECTURE.md",
    ROOT / "docs/DATABASE.md",
    ROOT / "docs/plans/2026-08-06-mvp-v1.1-development-plan.md",
    ROOT / "docs/plans/2026-08-06-mvp-v1.2-sprint3-life-story-engine.md",
    ROOT / "docs/plans/2026-08-06-auth-system.md",
]

ER_ASCII = """
┌─────────────────────────────────────────────────────────────────────────┐
│                        念念年年 · 数据库 ER 关系                          │
└─────────────────────────────────────────────────────────────────────────┘

  users ──────< family_users >────── families
    │                                  │
    │                                  ├──< photos ────< tags
    │                                  │       │
    │                                  │       └────|| memory_cards (1:1)
    │                                  │
    │                                  ├──< stories ──< story_memory_cards
    │                                  │       │              story_versions
    │                                  │       └──< shares
    │                                  │
    │                                  └──< life_movies ──< movie_chapters
    │                                              │              │
    │                                              └── movie_shares │
    │                                                               │
    └──── invitations                                    movie_chapters >── stories

  photos ──< photo_shares          verify_codes (OTP, 独立)
  photos ──< global_memory_search (检索读模型)

  共 17 张表 · SQLite · data/niannian.db
"""

DATA_FLOW_ASCII = """
  photos ──AI解析──> memory_cards + tags
       │
       └──同步──> global_memory_search
       
  memory_cards ──聚类──> stories + story_memory_cards + story_versions
       │
  stories ──编排──> life_movies + movie_chapters
       │
  life_movies ──MeloTTS──> 旁白 WAV ──FFmpeg──> MP4 (media_url)
"""


def set_doc_font(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("念念年年\n开发文档全集")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0xD9, 0x8A, 0x45)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "版本：2026-08-10\n"
        "仓库：github.com/SuDo-SmallLiu/niannian\n"
        "生产：niannian-years.top\n"
        "整理：cker + Cursor"
    )
    m.font.size = Pt(11)
    m.font.color.rgb = RGBColor(0x8B, 0x73, 0x55)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    doc.add_heading("目录", level=1)
    items = [
        "第一部分  项目概览与索引",
        "第二部分  产品路线图（2026-08-10）",
        "第三部分  开发历程与 Git 时间线",
        "第四部分  系统架构",
        "第五部分  数据库架构（17 张表）",
        "第六部分  MVP V1.1 开发计划（归档）",
        "第七部分  Sprint 3 Life Story Engine（归档）",
        "第八部分  Auth 系统设计（归档）",
        "附录 A    数据库 ER 关系图（文本版）",
        "附录 B    数据流图（文本版）",
    ]
    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}", style="List Number")
    doc.add_page_break()


def is_table_row(line: str) -> bool:
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and "|" in s[1:-1]


def parse_table_rows(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and is_table_row(lines[i]):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if all(re.match(r"^[-:\s]+$", c) for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            text = row[ci] if ci < len(row) else ""
            table.rows[ri].cells[ci].text = text.replace("`", "")


def render_markdown(doc: Document, text: str, part_title: str) -> None:
    doc.add_heading(part_title, level=1)
    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            if in_code:
                block = "\n".join(code_buf)
                if "mermaid" in code_buf[:1] or "erDiagram" in block or "flowchart" in block:
                    if "erDiagram" in block:
                        p = doc.add_paragraph("（ER 关系图见附录 A）")
                        p.runs[0].italic = True
                    elif "flowchart" in block or "sequenceDiagram" in block:
                        p = doc.add_paragraph("（流程图见正文 ASCII 版或附录 B）")
                        p.runs[0].italic = True
                else:
                    p = doc.add_paragraph(block)
                    p.style = "No Spacing"
                    for run in p.runs:
                        run.font.name = "Consolas"
                        run.font.size = Pt(9)
                code_buf = []
                in_code = False
            else:
                in_code = True
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        if is_table_row(line):
            rows, i = parse_table_rows(lines, i)
            add_table(doc, rows)
            continue

        if line.strip().startswith("- [x]"):
            doc.add_paragraph("☑ " + line.strip()[5:].strip(), style="List Bullet")
            i += 1
            continue
        if line.strip().startswith("- [ ]"):
            doc.add_paragraph("☐ " + line.strip()[5:].strip(), style="List Bullet")
            i += 1
            continue
        if line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
            i += 1
            continue

        if re.match(r"^\d+\.\s", line.strip()):
            doc.add_paragraph(re.sub(r"^\d+\.\s", "", line.strip()), style="List Number")
            i += 1
            continue

        clean = line.strip()
        clean = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", clean)
        clean = clean.replace("`", "")
        if clean.startswith(">"):
            p = doc.add_paragraph(clean.lstrip("> ").strip())
            p.runs[0].italic = True
        elif clean == "---":
            pass
        else:
            doc.add_paragraph(clean)
        i += 1


def main() -> None:
    doc = Document()
    set_doc_font(doc)
    sections = doc.sections
    for sec in sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(2.8)
        sec.right_margin = Cm(2.8)

    add_title_page(doc)
    add_toc(doc)

    part_names = [
        "第一部分  项目概览与索引",
        "第二部分  产品路线图（2026-08-10）",
        "第三部分  开发历程与 Git 时间线",
        "第四部分  系统架构",
        "第五部分  数据库架构（17 张表）",
        "第六部分  MVP V1.1 开发计划（归档）",
        "第七部分  Sprint 3 Life Story Engine（归档）",
        "第八部分  Auth 系统设计（归档）",
    ]

    for path, part in zip(DOCS, part_names):
        if path.exists():
            render_markdown(doc, path.read_text(encoding="utf-8"), part)
            doc.add_page_break()

    doc.add_heading("附录 A  数据库 ER 关系图（文本版）", level=1)
    p = doc.add_paragraph(ER_ASCII.strip())
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    doc.add_heading("附录 B  数据流图（文本版）", level=1)
    p = doc.add_paragraph(DATA_FLOW_ASCII.strip())
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)

    doc.add_heading("附录 C  17 张表速查", level=1)
    tables = [
        ("用户与家庭", "users, verify_codes, family_users, invitations, families"),
        ("照片与记忆", "photos, memory_cards, tags"),
        ("故事", "stories, story_memory_cards, story_versions, shares"),
        ("人生电影", "life_movies, movie_chapters, movie_shares"),
        ("分享", "photo_shares"),
        ("检索", "global_memory_search"),
    ]
    t = doc.add_table(rows=len(tables) + 1, cols=2)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "分组"
    t.rows[0].cells[1].text = "表名"
    for ri, (g, names) in enumerate(tables, 1):
        t.rows[ri].cells[0].text = g
        t.rows[ri].cells[1].text = names

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"✓ 已生成: {OUT}")
    print(f"  大小: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
